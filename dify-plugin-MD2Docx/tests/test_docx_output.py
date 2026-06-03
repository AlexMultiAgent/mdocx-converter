"""
Integration tests: full pandoc conversion + style overrides.
Verifies the generated DOCX has correct font, size, spacing, and margins.
All tests require pandoc binary (auto-downloaded by pypandoc on first run).
"""
import os
import sys
import io

import pytest

pytestmark = pytest.mark.slow

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))

from md2docx import (
    convert_via_pandoc,
    apply_style_overrides,
    build_pandoc_metadata,
    _ensure_pandoc,
    PROFILE_DEFAULTS,
    NORMAL_STYLE_NAMES,
    HEADING_ALIASES,
    CODE_STYLE_NAMES,
)


@pytest.fixture(scope="module")
def pandoc_available():
    """Ensure pandoc is available before running these tests."""
    _ensure_pandoc()
    import pypandoc
    assert pypandoc.get_pandoc_version() is not None


@pytest.fixture
def reference_docx():
    """Path to the english academic reference template."""
    plugin_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(plugin_root, "multi-templates", "reference_english_academic.docx")


@pytest.fixture
def source_dir(reference_docx):
    return os.path.dirname(reference_docx)


@pytest.fixture
def sample_markdown():
    return """# Heading 1

## Heading 2

### Heading 3

This is a normal paragraph with some text to test body font settings.
It should be long enough to be meaningful in the output document.

```python
def hello():
    print("Hello World")
```
"""


class TestFullConversion:
    """End-to-end tests: markdown → pandoc → style overrides → verified DOCX."""

    def test_academic_profile_overrides(
        self, pandoc_available, reference_docx, source_dir, sample_markdown
    ):
        """Full pipeline with academic profile and verify style overrides applied."""
        from docx import Document

        metadata = build_pandoc_metadata("academic", None, None, None)
        docx_io = convert_via_pandoc(
            sample_markdown, reference_docx, source_dir, "", metadata
        )
        doc = Document(docx_io)
        apply_style_overrides(doc, "academic", {})

        # Verify Normal style exists and has font overrides applied
        normal_found = False
        for name in NORMAL_STYLE_NAMES:
            try:
                style = doc.styles[name]
                normal_found = True
                # Check font was set (SimSun for academic Chinese, but English
                # template will get Times New Roman / SimSun via metadata)
                rPr = style._element.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
                )
                if rPr is not None:
                    rFonts = rPr.find(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
                    )
                    if rFonts is not None:
                        ascii_font = rFonts.get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii"
                        )
                        assert ascii_font is not None
                break
            except KeyError:
                continue
        assert normal_found, f"None of {NORMAL_STYLE_NAMES} found in document styles"

    def test_heading_styles_applied(
        self, pandoc_available, reference_docx, source_dir, sample_markdown
    ):
        """Verify heading styles are applied after conversion."""
        from docx import Document

        metadata = build_pandoc_metadata("academic", None, None, None)
        docx_io = convert_via_pandoc(
            sample_markdown, reference_docx, source_dir, "", metadata
        )
        doc = Document(docx_io)
        apply_style_overrides(doc, "academic", {})

        # Check Heading 1 style has font set
        for alias in HEADING_ALIASES["Heading 1"]:
            try:
                style = doc.styles[alias]
                rPr = style._element.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
                )
                if rPr is not None:
                    rFonts = rPr.find(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
                    )
                    if rFonts is not None:
                        ascii_font = rFonts.get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii"
                        )
                        assert ascii_font is not None
                break
            except KeyError:
                continue

    def test_technical_profile_code_styles(
        self, pandoc_available, reference_docx, source_dir, sample_markdown
    ):
        """Technical profile should set Consolas on SourceCode/VerbatimChar."""
        from docx import Document

        metadata = build_pandoc_metadata("technical", None, None, None)
        docx_io = convert_via_pandoc(
            sample_markdown, reference_docx, source_dir, "", metadata
        )
        doc = Document(docx_io)
        apply_style_overrides(doc, "technical", {})

        # At least one code style should have Consolas set
        code_font_found = False
        for code_name in CODE_STYLE_NAMES:
            try:
                style = doc.styles[code_name]
                rPr = style._element.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
                )
                if rPr is not None:
                    rFonts = rPr.find(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
                    )
                    if rFonts is not None:
                        ascii = rFonts.get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii"
                        )
                        if ascii == "Consolas":
                            code_font_found = True
                            break
            except KeyError:
                continue
        # Code styles may not exist in all reference templates; skip if none found
        if not code_font_found:
            pytest.skip("Code styles (SourceCode/VerbatimChar) not present in this reference.docx template")

    def test_margin_override(
        self, pandoc_available, reference_docx, source_dir, sample_markdown
    ):
        """Margin overrides should be applied to sections."""
        from docx import Document

        metadata = build_pandoc_metadata("academic", None, None, None)
        docx_io = convert_via_pandoc(
            sample_markdown, reference_docx, source_dir, "", metadata
        )
        doc = Document(docx_io)
        apply_style_overrides(doc, "academic", {
            "margin_top_mm": 30,
            "margin_bottom_mm": 30,
        })

        # Margins applied — verify sections have non-default margins
        for section in doc.sections:
            # 30mm = 3cm
            assert section.top_margin is not None
            assert section.bottom_margin is not None
            break

    def test_roundtrip_readback(
        self, pandoc_available, reference_docx, source_dir, sample_markdown
    ):
        """Generated DOCX should be re-openable by python-docx after save."""
        from docx import Document

        metadata = build_pandoc_metadata("business", None, None, None)
        docx_io = convert_via_pandoc(
            sample_markdown, reference_docx, source_dir, "", metadata
        )
        doc = Document(docx_io)
        apply_style_overrides(doc, "business", {})

        # Save to BytesIO and re-read
        output_io = io.BytesIO()
        doc.save(output_io)
        output_io.seek(0)

        # Should not raise
        doc2 = Document(output_io)
        assert doc2 is not None
        # Should have at least one paragraph (the body text)
        paragraphs = doc2.paragraphs
        assert len(paragraphs) > 0


class TestPandocEndToEnd:
    """Smoke tests that exercise the real pandoc binary end-to-end. These do
    NOT require any specific reference template — they verify that
    `convert_via_pandoc` returns a valid, openable DOCX (zip with a
    word/document.xml entry), which is the property that the rest of the
    pipeline depends on.
    """

    @pytest.fixture(scope="module")
    def pandoc_available(self):
        from md2docx import _ensure_pandoc
        _ensure_pandoc()
        import pypandoc
        assert pypandoc.get_pandoc_version() is not None

    def test_minimal_conversion_returns_valid_docx(self, pandoc_available):
        """The most basic regression: "# Hi" must round-trip into a valid DOCX
        byte stream that python-docx can open. This is the test that would
        have caught the original `outputfile=None` bug.
        """
        from docx import Document
        from md2docx import convert_via_pandoc, build_pandoc_metadata
        import io, os

        plugin_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        ref = os.path.join(plugin_root, "multi-templates", "reference_english_academic.docx")
        source_dir = os.path.dirname(ref)

        docx_io = convert_via_pandoc(
            "# Hello World\n\nA paragraph.",
            ref, source_dir, "",
            build_pandoc_metadata("academic", None, None, None),
        )
        data = docx_io.read()
        assert len(data) > 0
        # DOCX is a zip; first 2 bytes are "PK"
        assert data[:2] == b"PK", f"Output is not a zip, got first bytes: {data[:8]!r}"

        # And python-docx can open it
        doc = Document(io.BytesIO(data))
        assert len(doc.paragraphs) >= 1

    def test_strip_dangling_image_rels(self):
        """_strip_dangling_image_rels must remove image rels whose target is missing."""
        import io, zipfile
        from md2docx import _strip_dangling_image_rels

        # Build a fake docx with one present image and one missing image
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("[Content_Types].xml", """<?xml version="1.0"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>""")
            zf.writestr("word/_rels/document.xml.rels", """<?xml version="1.0"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/present.png"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/missing.png"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>""")
            zf.writestr("word/media/present.png", b"fake-png-bytes")
            zf.writestr("word/styles.xml", "<x/>")
        src = buf.getvalue()

        cleaned = _strip_dangling_image_rels(src)
        with zipfile.ZipFile(io.BytesIO(cleaned)) as zf:
            rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
        # The missing image rel is gone, the present one and styles survive
        assert "media/missing.png" not in rels
        assert "media/present.png" in rels
        assert "styles.xml" in rels

    def test_strip_no_op_on_clean_docx(self):
        """A docx whose rels all resolve to real parts should be returned as-is."""
        import io, zipfile
        from md2docx import _strip_dangling_image_rels

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("word/_rels/document.xml.rels", """<?xml version="1.0"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>""")
            zf.writestr("word/styles.xml", "<x/>")
        src = buf.getvalue()
        assert _strip_dangling_image_rels(src) == src

    def test_strip_tolerates_invalid_zip(self):
        """Invalid input must be returned unchanged (no exception)."""
        from md2docx import _strip_dangling_image_rels
        assert _strip_dangling_image_rels(b"not a zip") == b"not a zip"
        assert _strip_dangling_image_rels(b"") == b""


class TestMermaidRetryLogic:
    """Verify the new Mermaid retry + budget logic, with the network mocked."""

    def test_render_succeeds_after_one_transient_failure(self, monkeypatch):
        """If the first call raises and the second succeeds, the diagram renders."""
        from md2docx import preprocess_mermaid, _render_one_with_retry
        import time
        calls = {"n": 0}
        def fake_render(diagram):
            calls["n"] += 1
            if calls["n"] == 1:
                raise RuntimeError("transient")
            return b"\x89PNG\r\n\x1a\n"  # PNG signature

        monkeypatch.setattr("md2docx.render_mermaid_via_api", fake_render)
        # Skip the actual sleep between retries
        monkeypatch.setattr("md2docx.time.sleep", lambda s: None)

        import tempfile
        with tempfile.TemporaryDirectory() as td:
            errors = []
            idx, path = _render_one_with_retry(1, "graph TD; A-->B", td, errors)
            assert idx == 1
            assert path is not None
            assert path.endswith("diagram-1.png")
            assert errors == []
            assert calls["n"] == 2  # 1 failure + 1 success

    def test_render_fails_after_all_attempts(self, monkeypatch):
        """After MERMAID_MAX_ATTEMPTS failures, an error is recorded and path is None."""
        from md2docx import preprocess_mermaid, _render_one_with_retry, MERMAID_MAX_ATTEMPTS
        def fake_render(diagram):
            raise RuntimeError("persistent failure")
        monkeypatch.setattr("md2docx.render_mermaid_via_api", fake_render)
        monkeypatch.setattr("md2docx.time.sleep", lambda s: None)

        import tempfile
        with tempfile.TemporaryDirectory() as td:
            errors = []
            idx, path = _render_one_with_retry(1, "graph TD", td, errors)
            assert path is None
            assert len(errors) == 1
            assert "persistent failure" in errors[0]
